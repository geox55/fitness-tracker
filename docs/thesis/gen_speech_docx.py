"""Генератор docx-версии речи защиты из speech-maria.md.

Зачем отдельный скрипт, а не pandoc: держим оформление под контролем и
воспроизводимо (как build_thesis_docx.py) — заголовки, жирные ключевые цифры,
таблица-шпаргалка, блок вопросов. Источник — Markdown, чтобы правка текста была
в одном месте.

Запуск: `uv run python docs/thesis/gen_speech_docx.py`
Результат: docs/thesis/Речь-Лапова-Мария.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).parent
SRC = HERE / "speech-maria.md"
OUT = HERE / "Речь-Лапова-Мария.docx"

# Жирный фрагмент **...** разбивается на runs, чтобы выделить ключевые цифры.
_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_runs(paragraph, text: str) -> None:
    """Раскладывает inline-`**жирный**` на отдельные runs."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _emit_table(doc: Document, rows: list[list[str]]) -> None:
    """Markdown-таблица → docx-таблица (первая строка — шапка)."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, cols in enumerate(rows):
        for c, cell_text in enumerate(cols):
            cell = table.cell(r, c)
            cell.text = ""
            para = cell.paragraphs[0]
            _add_runs(para, cell_text)
            if r == 0:
                for run in para.runs:
                    run.bold = True


def build() -> Path:
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    # Базовый шрифт покрупнее — речь читают с листа.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    table_buf: list[list[str]] = []

    def flush_table() -> None:
        if table_buf:
            _emit_table(doc, table_buf)
            table_buf.clear()

    for raw in md:
        line = raw.rstrip()

        # Накопление строк markdown-таблицы.
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            # разделительная строка |---|---| — пропускаем
            if all(set(c) <= {"-", ":", " "} for c in cells):
                continue
            table_buf.append(cells)
            continue
        flush_table()

        if not line or line == "---":
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip().strip("«»"), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("> "):
            # Цитата-ремарка — серым курсивом.
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    flush_table()
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"Wrote: {out}")
